import base64
import csv
import html
import io
import re
from typing import Dict, List, Optional, Tuple

import pandas as pd
import streamlit as st

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Inches, Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


st.set_page_config(page_title="TriNetX Table 1 Generator", layout="wide")
st.title("TriNetX Baseline Patient Characteristics → Journal-Style Table 1")


REQUIRED_COLUMNS = [
    "Characteristic ID",
    "Characteristic Name",
    "Category",
    "Cohort 1 Before: Patient Count",
    "Cohort 1 Before: % of Cohort",
    "Cohort 2 Before: Patient Count",
    "Cohort 2 Before: % of Cohort",
    "Before: Standardized Mean Difference",
    "Cohort 1 After: Patient Count",
    "Cohort 1 After: % of Cohort",
    "Cohort 2 After: Patient Count",
    "Cohort 2 After: % of Cohort",
    "After: Standardized Mean Difference",
]

DEMOGRAPHIC_IDS = {"AI", "M", "F", "2106-3", "2054-5", "2135-2", "2028-9"}
DEMOGRAPHIC_ORDER = {
    "AI": 10,
    "M": 20,
    "F": 30,
    "2106-3": 40,  # White
    "2054-5": 50,  # Black or African American
    "2135-2": 60,  # Hispanic or Latino
    "2028-9": 70,  # Asian
}

LAB_ORDER = {
    "9083": 10,  # BMI
    "9002": 20,  # LDL Cholesterol
    "9004": 30,  # Triglycerides
}

MEDICATION_ORDER = {
    "CV300": 10,  # Antiarrhythmics
    "CV100": 20,  # Beta Blockers
    "CV700": 30,  # Diuretics
    "CV800": 40,  # ACE Inhibitors
    "CV200": 50,  # Calcium Channel Blockers
    "CV805": 60,  # Angiotensin II Inhibitors
}

SECTION_ORDER = {
    "Demographics": 1,
    "Diagnoses": 2,
    "Labs": 3,
    "Medications": 4,
    "Other": 5,
}

SECTION_IDENTIFIER_LABEL = {
    "Demographics": "Identifier Code",
    "Diagnoses": "ICD-10",
    "Labs": "LOINC",
    "Medications": "",
    "Other": "Identifier Code",
}

NAME_REPLACEMENTS = {
    "ANTIARRHYTHMICS": "Antiarrhythmics",
    "BETA BLOCKERS/RELATED": "Beta Blockers",
    "DIURETICS": "Diuretics",
    "ACE INHIBITORS": "ACE Inhibitors",
    "CALCIUM CHANNEL BLOCKERS": "Calcium Channel Blockers",
    "ANGIOTENSIN II INHIBITOR": "Angiotensin II Inhibitors",
    "Triglyceride [Mass/volume] in Serum, Plasma or Blood": "Triglycerides",
    "Cholesterol in LDL [Mass/volume] in Serum or Plasma": "LDL Cholesterol",
}


def decode_uploaded_file(uploaded_file) -> str:
    """Decode a TriNetX CSV upload, tolerating common encodings."""
    raw_bytes = uploaded_file.getvalue()
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="replace")


def find_header_row(lines: List[str]) -> Optional[int]:
    """Find the TriNetX header row even when the export contains title/note rows."""
    for i, line in enumerate(lines[:75]):
        try:
            cols = [c.strip() for c in next(csv.reader([line]))]
        except Exception:
            continue
        if {"Characteristic ID", "Characteristic Name", "Category"}.issubset(set(cols)):
            return i
    return None


def read_trinetx_baseline_csv(uploaded_file) -> pd.DataFrame:
    """Read a Baseline Patient Characteristics CSV exported by TriNetX."""
    content = decode_uploaded_file(uploaded_file)
    lines = content.splitlines()
    header_row = find_header_row(lines)

    if header_row is None:
        raise ValueError(
            "Could not find the TriNetX header row. The file must contain "
            "'Characteristic ID', 'Characteristic Name', and 'Category'."
        )

    df = pd.read_csv(io.StringIO("\n".join(lines[header_row:])))
    df.columns = [str(c).strip() for c in df.columns]
    df = df.dropna(how="all").reset_index(drop=True)

    missing = [c for c in REQUIRED_COLUMNS if c not in df.columns]
    if missing:
        raise ValueError("The file is missing required columns: " + ", ".join(missing))

    for col in df.columns:
        if col not in ["Characteristic ID", "Characteristic Name", "Category"]:
            df[col] = pd.to_numeric(
                df[col].astype(str).str.replace("%", "", regex=False).str.replace(",", "", regex=False),
                errors="coerce",
            )
    return df


def safe_str(value) -> str:
    if pd.isna(value):
        return ""
    return str(value).strip()


def as_float(value) -> Optional[float]:
    if value is None or pd.isna(value):
        return None
    try:
        if isinstance(value, str):
            value = value.replace("%", "").replace(",", "").strip()
            if not value:
                return None
        return float(value)
    except Exception:
        return None


def format_count(value) -> str:
    number = as_float(value)
    if number is None:
        return ""
    return f"{int(round(number)):,}"


def format_percent(value, decimals: int = 2) -> str:
    number = as_float(value)
    if number is None:
        return ""
    return f"{number:.{decimals}f}%"


def format_count_percent(row: pd.Series, phase: str, cohort: int, pct_decimals: int) -> str:
    count = format_count(row.get(f"Cohort {cohort} {phase}: Patient Count"))
    pct = format_percent(row.get(f"Cohort {cohort} {phase}: % of Cohort"), pct_decimals)
    if not count and not pct:
        return ""
    if count and pct:
        return f"{count} ({pct})"
    return count or pct


def format_mean_sd(row: pd.Series, phase: str, cohort: int, mean_decimals: int) -> str:
    mean = as_float(row.get(f"Cohort {cohort} {phase}: Mean"))
    sd = as_float(row.get(f"Cohort {cohort} {phase}: SD"))
    if mean is None or sd is None:
        return ""
    return f"{mean:.{mean_decimals}f} (±{sd:.{mean_decimals}f})"


def format_smd(value, decimals: int, dynamic: bool = True) -> str:
    number = as_float(value)
    if number is None:
        return ""
    if dynamic:
        # Matches the usual Table 1 convention: larger imbalances are easier to scan,
        # very small post-match imbalances retain three decimals.
        decimals = 2 if abs(number) >= 0.10 else 3
    return f"{number:.{decimals}f}"


def format_p_value(value) -> str:
    number = as_float(value)
    if number is None:
        return ""
    if number == 0 or number < 0.001:
        return "p<.001"
    return f"{number:.3f}"


def has_mean_sd(row: pd.Series) -> bool:
    needed = [
        "Cohort 1 Before: Mean",
        "Cohort 1 Before: SD",
        "Cohort 2 Before: Mean",
        "Cohort 2 Before: SD",
        "Cohort 1 After: Mean",
        "Cohort 1 After: SD",
        "Cohort 2 After: Mean",
        "Cohort 2 After: SD",
    ]
    return any(as_float(row.get(c)) is not None for c in needed)


def clean_name(name: str, clean_labels: bool = True) -> str:
    name = safe_str(name)
    if not clean_labels:
        return name

    if name in NAME_REPLACEMENTS:
        return NAME_REPLACEMENTS[name]

    if name.isupper() and any(ch.isalpha() for ch in name):
        name = name.title()
        name = (
            name.replace("Ace ", "ACE ")
            .replace("Ii ", "II ")
            .replace("Ldl", "LDL")
            .replace("Bmi", "BMI")
        )
    return name


def clean_category(category: str, simplify_open_ended_lab_bins: bool = True) -> str:
    category = safe_str(category)
    if not category:
        return ""

    if simplify_open_ended_lab_bins:
        # TriNetX often exports high lab categories as "151-500 mg/dL" or "101-500 mg/dL";
        # journal tables are typically clearer when the upper cap is rendered as "+".
        category = re.sub(r"^(\d+)-500\s+(mg/dL)$", r"\1+ \2", category)

    return category.replace("m2", "m²")


def infer_section(row: pd.Series) -> str:
    code = safe_str(row.get("Characteristic ID"))
    name = safe_str(row.get("Characteristic Name"))

    if code in DEMOGRAPHIC_IDS:
        return "Demographics"
    if code.upper().startswith("CV"):
        return "Medications"
    if re.fullmatch(r"\d+", code):
        return "Labs"
    if name.isupper() and code and not code[0].isdigit():
        return "Medications"
    if code:
        return "Diagnoses"
    return "Other"


def is_continuous_summary_row(row: pd.Series) -> bool:
    return safe_str(row.get("Category")) == "" and has_mean_sd(row)


def should_include_row(row: pd.Series, exclude_aggregate_lab_rows: bool) -> bool:
    section = infer_section(row)
    if is_continuous_summary_row(row):
        # Age at Index is part of the target Table 1. Aggregate lab rows are often excluded
        # when the table uses clinically meaningful bins instead.
        if safe_str(row.get("Characteristic ID")) == "AI":
            return True
        if exclude_aggregate_lab_rows and section == "Labs":
            return False
    return True


def make_display_label(
    row: pd.Series,
    clean_labels: bool,
    simplify_open_ended_lab_bins: bool,
) -> str:
    code = safe_str(row.get("Characteristic ID"))
    name = clean_name(row.get("Characteristic Name"), clean_labels)
    category = clean_category(row.get("Category"), simplify_open_ended_lab_bins)

    if code == "AI":
        return "Age at Index, mean (SD)"
    if category:
        return f"{name} ({category})"
    if is_continuous_summary_row(row):
        return f"{name}, mean (SD)"
    return name


def row_sort_key(row: pd.Series) -> Tuple[int, int, int]:
    section = infer_section(row)
    code = safe_str(row.get("Characteristic ID"))
    original_index = int(row.name) if row.name is not None else 9999

    if section == "Demographics":
        within_section_order = DEMOGRAPHIC_ORDER.get(code, original_index + 100)
    elif section == "Labs":
        within_section_order = LAB_ORDER.get(code, original_index + 100)
    elif section == "Medications":
        within_section_order = MEDICATION_ORDER.get(code.upper(), original_index + 100)
    else:
        within_section_order = original_index + 100

    return (
        SECTION_ORDER.get(section, 99),
        within_section_order,
        original_index,
    )


def build_publication_rows(
    raw_df: pd.DataFrame,
    cohort_1_label: str,
    cohort_2_label: str,
    pct_decimals: int,
    mean_decimals: int,
    smd_decimals: int,
    dynamic_smd_decimals: bool,
    include_p_values: bool,
    clean_labels: bool,
    simplify_open_ended_lab_bins: bool,
    exclude_aggregate_lab_rows: bool,
    blank_repeated_lab_codes: bool,
) -> pd.DataFrame:
    rows = []

    working_df = raw_df.copy()
    working_df["_section"] = working_df.apply(infer_section, axis=1)
    working_df["_include"] = working_df.apply(
        lambda r: should_include_row(r, exclude_aggregate_lab_rows),
        axis=1,
    )
    working_df = working_df[working_df["_include"]].copy()
    working_df["_sort"] = working_df.apply(row_sort_key, axis=1)
    working_df = working_df.sort_values("_sort", kind="stable")

    for section, section_df in working_df.groupby("_section", sort=False):
        rows.append(
            {
                "Include": True,
                "Order": len(rows) + 1,
                "Row Type": "group",
                "Section": section,
                "Characteristic": section,
                "Identifier Code": SECTION_IDENTIFIER_LABEL.get(section, "Identifier Code"),
                "Before: " + cohort_1_label: "",
                "Before: " + cohort_2_label: "",
                "Before: SMD": "",
                "After: " + cohort_1_label: "",
                "After: " + cohort_2_label: "",
                "After: SMD": "",
                "Before: p-Value": "",
                "After: p-Value": "",
            }
        )

        previous_lab_code = None
        for _, raw_row in section_df.iterrows():
            continuous = is_continuous_summary_row(raw_row)
            label = make_display_label(raw_row, clean_labels, simplify_open_ended_lab_bins)
            code = safe_str(raw_row.get("Characteristic ID"))

            if section == "Labs" and blank_repeated_lab_codes:
                display_code = "" if code == previous_lab_code else code
                previous_lab_code = code
            else:
                display_code = code

            if continuous:
                before_c1 = format_mean_sd(raw_row, "Before", 1, mean_decimals)
                before_c2 = format_mean_sd(raw_row, "Before", 2, mean_decimals)
                after_c1 = format_mean_sd(raw_row, "After", 1, mean_decimals)
                after_c2 = format_mean_sd(raw_row, "After", 2, mean_decimals)
            else:
                before_c1 = format_count_percent(raw_row, "Before", 1, pct_decimals)
                before_c2 = format_count_percent(raw_row, "Before", 2, pct_decimals)
                after_c1 = format_count_percent(raw_row, "After", 1, pct_decimals)
                after_c2 = format_count_percent(raw_row, "After", 2, pct_decimals)

            rows.append(
                {
                    "Include": True,
                    "Order": len(rows) + 1,
                    "Row Type": "data",
                    "Section": section,
                    "Characteristic": label,
                    "Identifier Code": display_code,
                    "Before: " + cohort_1_label: before_c1,
                    "Before: " + cohort_2_label: before_c2,
                    "Before: SMD": format_smd(
                        raw_row.get("Before: Standardized Mean Difference"),
                        smd_decimals,
                        dynamic_smd_decimals,
                    ),
                    "After: " + cohort_1_label: after_c1,
                    "After: " + cohort_2_label: after_c2,
                    "After: SMD": format_smd(
                        raw_row.get("After: Standardized Mean Difference"),
                        smd_decimals,
                        dynamic_smd_decimals,
                    ),
                    "Before: p-Value": format_p_value(raw_row.get("Before: p-Value")),
                    "After: p-Value": format_p_value(raw_row.get("After: p-Value")),
                }
            )

    table_df = pd.DataFrame(rows)

    if not include_p_values:
        table_df = table_df.drop(columns=["Before: p-Value", "After: p-Value"], errors="ignore")

    return table_df


def display_columns(table_df: pd.DataFrame) -> List[str]:
    hidden = {"Include", "Order", "Row Type", "Section"}
    return [c for c in table_df.columns if c not in hidden]


def make_html_table(
    table_df: pd.DataFrame,
    table_title: str,
    cohort_1_label: str,
    cohort_2_label: str,
    font_size: int,
    include_p_values: bool,
    table_width_percent: int,
) -> str:
    before_span = 4 if include_p_values else 3
    after_span = 4 if include_p_values else 3

    before_headers = [
        cohort_1_label + " (%)",
        cohort_2_label + " (%)",
        "SMD",
    ]
    after_headers = [
        cohort_1_label + " (%)",
        cohort_2_label + " (%)",
        "SMD",
    ]

    if include_p_values:
        before_headers.append("p-Value")
        after_headers.append("p-Value")

    css = f"""
<style>
.table1-wrap {{
    width: {table_width_percent}%;
    overflow-x: auto;
}}
.table1 {{
    border-collapse: collapse;
    width: 100%;
    font-family: Arial, Helvetica, sans-serif;
    font-size: {font_size}pt;
    line-height: 1.18;
}}
.table1 caption {{
    caption-side: top;
    text-align: left;
    font-weight: bold;
    margin-bottom: 8px;
}}
.table1 th, .table1 td {{
    border: 1px solid #222;
    padding: 5px 7px;
    vertical-align: middle;
}}
.table1 th {{
    background: #f2f2f2;
    font-weight: bold;
    text-align: center;
}}
.table1 td:first-child {{
    text-align: left;
    min-width: 220px;
}}
.table1 td:nth-child(2) {{
    text-align: center;
    min-width: 82px;
}}
.table1 td:not(:first-child):not(:nth-child(2)) {{
    text-align: center;
    white-space: nowrap;
}}
.table1 .group-row td {{
    background: #e8e8e8;
    font-weight: bold;
}}
.table1 .group-row td:first-child {{
    text-align: left;
}}
</style>
"""

    html_parts = [css, '<div class="table1-wrap">', '<table class="table1">']
    if table_title:
        html_parts.append(f"<caption>{html.escape(table_title)}</caption>")

    html_parts.append(
        "<thead>"
        "<tr>"
        "<th rowspan='2'></th>"
        "<th rowspan='2'></th>"
        f"<th colspan='{before_span}'>Before Propensity Score Matching</th>"
        f"<th colspan='{after_span}'>After Propensity Score Matching</th>"
        "</tr>"
    )

    html_parts.append("<tr>")
    for h in before_headers + after_headers:
        html_parts.append(f"<th>{html.escape(h)}</th>")
    html_parts.append("</tr></thead><tbody>")

    for _, row in table_df.iterrows():
        if not bool(row.get("Include", True)):
            continue

        row_type = safe_str(row.get("Row Type"))
        values = []
        if row_type == "group":
            values = [
                row.get("Characteristic", ""),
                row.get("Identifier Code", ""),
                "",
                "",
                "",
            ]
            if include_p_values:
                values.append("")
            values += ["", "", ""]
            if include_p_values:
                values.append("")
            html_parts.append("<tr class='group-row'>")
        else:
            values = [
                row.get("Characteristic", ""),
                row.get("Identifier Code", ""),
                row.get("Before: " + cohort_1_label, ""),
                row.get("Before: " + cohort_2_label, ""),
                row.get("Before: SMD", ""),
            ]
            if include_p_values:
                values.append(row.get("Before: p-Value", ""))
            values += [
                row.get("After: " + cohort_1_label, ""),
                row.get("After: " + cohort_2_label, ""),
                row.get("After: SMD", ""),
            ]
            if include_p_values:
                values.append(row.get("After: p-Value", ""))
            html_parts.append("<tr>")

        for cell in values:
            html_parts.append(f"<td>{html.escape(safe_str(cell))}</td>")
        html_parts.append("</tr>")

    html_parts.append("</tbody></table></div>")
    return "".join(html_parts)


def make_plain_export_df(
    table_df: pd.DataFrame,
    cohort_1_label: str,
    cohort_2_label: str,
    include_p_values: bool,
) -> pd.DataFrame:
    rows = []
    for _, row in table_df.iterrows():
        if not bool(row.get("Include", True)):
            continue

        export_row = {
            "Characteristic": row.get("Characteristic", ""),
            "Identifier Code": row.get("Identifier Code", ""),
            "Before Propensity Score Matching - " + cohort_1_label: row.get("Before: " + cohort_1_label, ""),
            "Before Propensity Score Matching - " + cohort_2_label: row.get("Before: " + cohort_2_label, ""),
            "Before Propensity Score Matching - SMD": row.get("Before: SMD", ""),
            "After Propensity Score Matching - " + cohort_1_label: row.get("After: " + cohort_1_label, ""),
            "After Propensity Score Matching - " + cohort_2_label: row.get("After: " + cohort_2_label, ""),
            "After Propensity Score Matching - SMD": row.get("After: SMD", ""),
        }
        if include_p_values:
            export_row["Before Propensity Score Matching - p-Value"] = row.get("Before: p-Value", "")
            export_row["After Propensity Score Matching - p-Value"] = row.get("After: p-Value", "")
        rows.append(export_row)

    return pd.DataFrame(rows)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 8):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(safe_str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def make_docx_bytes(
    table_df: pd.DataFrame,
    table_title: str,
    cohort_1_label: str,
    cohort_2_label: str,
    include_p_values: bool,
    font_size: int,
) -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed.")

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    if table_title:
        p = document.add_paragraph()
        run = p.add_run(table_title)
        run.bold = True
        run.font.size = Pt(10)

    before_span = 4 if include_p_values else 3
    after_span = 4 if include_p_values else 3
    n_cols = 2 + before_span + after_span

    display_df = table_df[table_df["Include"].astype(bool)].copy()

    table = document.add_table(rows=2 + len(display_df), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row 1
    set_cell_text(table.cell(0, 0), "", bold=True, font_size=font_size)
    set_cell_text(table.cell(0, 1), "", bold=True, font_size=font_size)
    before_cell = table.cell(0, 2)
    before_cell.merge(table.cell(0, 2 + before_span - 1))
    set_cell_text(before_cell, "Before Propensity Score Matching", bold=True, font_size=font_size)

    after_cell = table.cell(0, 2 + before_span)
    after_cell.merge(table.cell(0, n_cols - 1))
    set_cell_text(after_cell, "After Propensity Score Matching", bold=True, font_size=font_size)

    # Header row 2
    headers = ["", ""]
    headers += [cohort_1_label + " (%)", cohort_2_label + " (%)", "SMD"]
    if include_p_values:
        headers.append("p-Value")
    headers += [cohort_1_label + " (%)", cohort_2_label + " (%)", "SMD"]
    if include_p_values:
        headers.append("p-Value")

    for j, header in enumerate(headers):
        set_cell_text(table.cell(1, j), header, bold=True, font_size=font_size)
        set_cell_shading(table.cell(1, j), "F2F2F2")

    # Data rows
    for i, (_, row) in enumerate(display_df.iterrows(), start=2):
        row_type = safe_str(row.get("Row Type"))
        if row_type == "group":
            values = [
                row.get("Characteristic", ""),
                row.get("Identifier Code", ""),
                "",
                "",
                "",
            ]
            if include_p_values:
                values.append("")
            values += ["", "", ""]
            if include_p_values:
                values.append("")
            bold = True
            fill = "E8E8E8"
        else:
            values = [
                row.get("Characteristic", ""),
                row.get("Identifier Code", ""),
                row.get("Before: " + cohort_1_label, ""),
                row.get("Before: " + cohort_2_label, ""),
                row.get("Before: SMD", ""),
            ]
            if include_p_values:
                values.append(row.get("Before: p-Value", ""))
            values += [
                row.get("After: " + cohort_1_label, ""),
                row.get("After: " + cohort_2_label, ""),
                row.get("After: SMD", ""),
            ]
            if include_p_values:
                values.append(row.get("After: p-Value", ""))
            bold = False
            fill = None

        for j, value in enumerate(values):
            cell = table.cell(i, j)
            set_cell_text(cell, value, bold=bold, font_size=font_size)
            if fill:
                set_cell_shading(cell, fill)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


with st.sidebar:
    st.header("Upload and table labels")
    uploaded_file = st.file_uploader("Upload TriNetX Baseline Patient Characteristics CSV", type=["csv"])

    table_title = st.text_input(
        "Table title",
        "Table 1: Main analysis cohort characteristics before and after propensity score matching.",
    )
    cohort_1_label = st.text_input("Cohort 1 label", "Statins")
    cohort_2_label = st.text_input("Cohort 2 label", "Control")

    st.header("Automatic formatting")
    clean_labels = st.checkbox("Clean TriNetX labels", value=True)
    exclude_aggregate_lab_rows = st.checkbox(
        "Exclude aggregate lab mean rows when lab bins are present",
        value=True,
        help="Keeps age as mean (SD), but hides aggregate BMI/LDL/triglyceride mean rows when binned lab rows are present.",
    )
    simplify_open_ended_lab_bins = st.checkbox(
        "Simplify high lab bins, e.g., 101-500 mg/dL → 101+ mg/dL",
        value=True,
    )
    blank_repeated_lab_codes = st.checkbox("Blank repeated LOINC codes within lab categories", value=True)
    include_p_values = st.checkbox("Include p-value columns", value=False)

    st.header("Number formatting")
    pct_decimals = st.slider("Percentage decimals", 0, 3, 2)
    mean_decimals = st.slider("Mean/SD decimals", 0, 3, 2)
    dynamic_smd_decimals = st.checkbox("Use dynamic SMD decimals", value=True)
    smd_decimals = st.slider("SMD decimals if dynamic is off", 2, 4, 3)

    st.header("Preview style")
    font_size = st.slider("Preview/DOCX font size", 6, 12, 8)
    table_width_percent = st.slider("Preview width (%)", 70, 100, 100)


if uploaded_file is None:
    st.info("Upload the Baseline Patient Characteristics CSV exported by TriNetX to generate the Table 1 layout.")
    st.stop()

try:
    raw_df = read_trinetx_baseline_csv(uploaded_file)
except Exception as exc:
    st.error(str(exc))
    st.stop()

st.success(f"Loaded {len(raw_df):,} baseline characteristic rows from the TriNetX export.")

publication_df = build_publication_rows(
    raw_df=raw_df,
    cohort_1_label=cohort_1_label,
    cohort_2_label=cohort_2_label,
    pct_decimals=pct_decimals,
    mean_decimals=mean_decimals,
    smd_decimals=smd_decimals,
    dynamic_smd_decimals=dynamic_smd_decimals,
    include_p_values=include_p_values,
    clean_labels=clean_labels,
    simplify_open_ended_lab_bins=simplify_open_ended_lab_bins,
    exclude_aggregate_lab_rows=exclude_aggregate_lab_rows,
    blank_repeated_lab_codes=blank_repeated_lab_codes,
)

st.subheader("Edit final table before export")
st.caption(
    "Use Include to hide rows, edit labels/values directly, and adjust Order to reorder rows. "
    "Group rows are retained so the exported table matches the Word-style Table 1 structure."
)

editor_cols = ["Include", "Order", "Row Type", "Section"] + display_columns(publication_df)
edited_df = st.data_editor(
    publication_df[editor_cols],
    use_container_width=True,
    height=520,
    disabled=["Row Type", "Section"],
    column_config={
        "Include": st.column_config.CheckboxColumn("Include"),
        "Order": st.column_config.NumberColumn("Order", min_value=1, step=1),
    },
)

edited_df = edited_df.copy()
edited_df["Order"] = pd.to_numeric(edited_df["Order"], errors="coerce").fillna(9999)
edited_df = edited_df.sort_values(["Order"], kind="stable").reset_index(drop=True)

st.subheader("Formatted Table 1 preview")
html_table = make_html_table(
    edited_df,
    table_title=table_title,
    cohort_1_label=cohort_1_label,
    cohort_2_label=cohort_2_label,
    font_size=font_size,
    include_p_values=include_p_values,
    table_width_percent=table_width_percent,
)
st.markdown(html_table, unsafe_allow_html=True)

plain_export_df = make_plain_export_df(
    edited_df,
    cohort_1_label=cohort_1_label,
    cohort_2_label=cohort_2_label,
    include_p_values=include_p_values,
)

csv_bytes = plain_export_df.to_csv(index=False).encode("utf-8")
html_bytes = html_table.encode("utf-8")

download_cols = st.columns(3)
with download_cols[0]:
    st.download_button(
        "Download table as CSV",
        data=csv_bytes,
        file_name="trinetx_table1_publication_ready.csv",
        mime="text/csv",
        use_container_width=True,
    )

with download_cols[1]:
    st.download_button(
        "Download table as HTML",
        data=html_bytes,
        file_name="trinetx_table1_publication_ready.html",
        mime="text/html",
        use_container_width=True,
    )

with download_cols[2]:
    if DOCX_AVAILABLE:
        try:
            docx_bytes = make_docx_bytes(
                edited_df,
                table_title=table_title,
                cohort_1_label=cohort_1_label,
                cohort_2_label=cohort_2_label,
                include_p_values=include_p_values,
                font_size=font_size,
            )
            st.download_button(
                "Download table as DOCX",
                data=docx_bytes,
                file_name="trinetx_table1_publication_ready.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as exc:
            st.warning(f"DOCX export failed: {exc}")
    else:
        st.warning("Install python-docx to enable DOCX export: pip install python-docx")

with st.expander("Copy/paste HTML"):
    st.code(html_table, language="html")

with st.expander("Raw TriNetX rows detected"):
    st.dataframe(raw_df, use_container_width=True)
